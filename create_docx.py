import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('SkillBridge-AI: Enterprise Agentic Workflow Orchestrator', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading('1. Executive Summary', level=1)
    
    doc.add_heading('The Challenge', level=2)
    doc.add_paragraph('ET Gen AI Hackathon Problem Statement 2:')
    p = doc.add_paragraph()
    p.add_run('"Design a multi-agent system that takes ownership of a complex, multi-step enterprise process. It should detect failures, self-correct, and complete the job with minimal human involvement — while keeping an auditable trail of every decision it makes."').italic = True
    
    doc.add_heading('The Solution: SkillBridge-AI', level=2)
    p = doc.add_paragraph()
    p.add_run('SkillBridge-AI').bold = True
    p.add_run(' completely owns and executes the complex, high-friction enterprise workflow of Technical Candidate Screening and Skill-Gap Analysis.')
    
    p = doc.add_paragraph()
    p.add_run('The Final Goal: ').bold = True
    p.add_run('Replace a slow, manual 45-minute HR screening task with a deterministic, self-correcting 3-second autonomous workflow. The system is designed to respect data privacy (100% offline-native, meaning no PII ever leaves the corporate firewall) while generating an immutable audit log mapping every programmatic decision made by the agents.')
    
    doc.add_heading('2. Multi-Agent Architecture (The Decision Graph)', level=1)
    doc.add_paragraph('Unlike a simple linear pipeline, SkillBridge-AI operates as an Intelligent Decision Graph. The system evaluates real-time data at every node and conditionally routes execution to the appropriate specialized sub-agent.')
    
    doc.add_heading('The Autonomous Agents', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('AgentOrchestrator: ').bold = True
    p.add_run('The brain of the operation. It assesses SLAs, monitors workflow health, and dictates payload routing.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('IngestionAgent & ExtractionAgent: ').bold = True
    p.add_run('NLP specialists responsible for isolating named entities and grammatical components without relying on static keyword lists.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('AnalysisAgent: ').bold = True
    p.add_run('Runs a proprietary 7-Pass Classifier backed by an Implication Engine (utilizing FAISS vectors) to map a candidate\'s latent skills to the enterprise\'s process requirements.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('StrategyAgent, FastTrackAgent, RedirectAgent: ').bold = True
    p.add_run('Specialized execution agents that synthesize actionable plans relying entirely on Orchestrator conditionals.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('ComplianceAgent: ').bold = True
    p.add_run('The final gatekeeper. Reviews all generated LLM tokens for bias, hallucinations, and format anomalies before delivery.')
    
    doc.add_heading('3. Fulfilling the Judging Criteria', level=1)
    
    doc.add_heading('A. Depth of Autonomy', level=2)
    doc.add_paragraph('The system achieves 100% autonomous completion without human involvement (unless explicitly escalated).')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Zero-Shot NLP Extraction: ').bold = True
    p.add_run('The Orchestrator does not utilize hardcoded lists of "skills to look for." It dynamically understands whether a word functions as an actionable technical ability in standard English syntax.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Intelligent Routing: ').bold = True
    p.add_run('Based on the AnalysisAgent\'s calculated vector-match score, the Orchestrator autonomously decides how to treat the candidate. High-scoring candidates are instantly routed to the FastTrackAgent to auto-generate interview questions; low-scoring candidates trigger the RedirectAgent to map their vectors to alternative enterprise roles.')
    
    doc.add_heading('B. Quality of Error Recovery (Self-Correction)', level=2)
    doc.add_paragraph('SkillBridge-AI is built to expect and handle failure gracefully:')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Extraction Failures: ').bold = True
    p.add_run('If an uploaded PDF provides 0 viable NLP skills, the Orchestrator flags a scrape failure and autonomously reroutes the document back to the IngestionAgent, activating an expensive but accurate optical-character-recognition (OCR) fallback routine.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Compliance Bounds Checking: ').bold = True
    p.add_run('If the StrategyAgent attempts to output biased or hallucinated upskilling advice, the ComplianceAgent fails the check. The Orchestrator immediately triggers a Self-Correction Loop, forcing the StrategyAgent to regenerate the plan with embedded constraints limits (up to 2 times).')
    
    doc.add_heading('C. Auditability', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Immutable Audit Ledger: ').bold = True
    p.add_run('The core requirement for enterprise deployment. The AuditLogger acts as the backbone of the entire workflow.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Every single processing step logs: 1. The specific agent running. 2. Sub-millisecond execution duration (vital for tracking SLA budgets). 3. Routing decisions with reasoning and confidence scores. 4. Alternatives considered at the node.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Complete visual transparency is surfaced on the UI Dashboard and exportable via JSON for HR compliance.')
    
    doc.add_heading('D. Real-World Applicability', level=2)
    doc.add_paragraph('Enterprise IT organizations fundamentally cannot pass unfiltered candidate PII (names, phone numbers, addresses) or proprietary job architectures to consumer LLMs like ChatGPT or Claude.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('100% Offline-Native Architecture: ').bold = True
    p.add_run('SkillBridge-AI relies entirely on local compute.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('spaCy handles NLP grammatical analysis.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('sentence-transformers handles embedding extraction.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('FAISS handles sub-millisecond vector similarity.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Ollama (Llama 3.2 1b/3b) handles generative text actions.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Zero APIs. Zero external network requests. Utterly air-gapped.')
    
    doc.add_heading('4. Technical Stack', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Domain'
    hdr_cells[1].text = 'Technology Layer'
    
    stacks = [
        ('User Interface', 'Streamlit (Python 3.10+)'),
        ('Agent Orchestration', 'Custom Python State-Machine & Routing Logic'),
        ('Vector Database', 'Meta FAISS (CPU-Optimized)'),
        ('NLP & Extraction', 'en_core_web_sm (spaCy), Regex Pattern Filtering'),
        ('Embedding Models', 'all-MiniLM-L6-v2 (sentence-transformers)'),
        ('Generative LLM (Local)', 'Ollama (Llama 3.2 1b/3b)')
    ]
    
    for domain, tech in stacks:
        row_cells = table.add_row().cells
        row_cells[0].text = domain
        row_cells[1].text = tech

    doc.add_heading('5. Running the Orchestrator', level=1)
    doc.add_paragraph('Prerequisites: Python 3.10+, Ollama.')
    p = doc.add_paragraph('ollama pull llama3.2\nollama serve\npip install -r requirements.txt\npython -m spacy download en_core_web_sm\nstreamlit run app.py')
    p.style = 'Macro Text'
    
    # Save word doc
    output_path = "Project_Documentation.docx"
    doc.save(output_path)
    print(f"Successfully generated {output_path}")

if __name__ == "__main__":
    main()
